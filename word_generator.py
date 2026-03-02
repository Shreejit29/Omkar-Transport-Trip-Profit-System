from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO


def add_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)

            tcPr.append(tcBorders)


def generate_word(trip_date, source, destination, material,
                  rate, weight, total_fare,
                  advance, pending_payment,
                  diesel, toll, food, driver,
                  other_expenses,
                  total_expenses, profit):

    doc = Document()
    doc.add_heading("Omkar Transport - Trip Report", level=1)

    # ---------------- Trip Details ----------------
    trip_data = [
        ["Sr. No.", "Particular", "Value"],
        ["1", "Date", str(trip_date)],
        ["2", "Source", source],
        ["3", "Destination", destination],
        ["4", "Material", material],
        ["5", "Rate per Tonne", f"₹ {rate:.2f}"],
        ["6", "Weight", f"{weight:.2f} Tonnes"],
        ["7", "Total Fare", f"₹ {total_fare:.2f}"],
    ]

    table = doc.add_table(rows=len(trip_data), cols=3)
    for i, row in enumerate(trip_data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
    add_table_borders(table)

    # ---------------- Payment Details ----------------
    doc.add_heading("Payment Details", level=2)

    payment_data = [
        ["Sr. No.", "Particular", "Amount"],
        ["1", "Advance Received", f"₹ {advance:.2f}"],
        ["2", "Pending Payment", f"₹ {pending_payment:.2f}"],
    ]

    table = doc.add_table(rows=len(payment_data), cols=3)
    for i, row in enumerate(payment_data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
    add_table_borders(table)

    # ---------------- Expense Details ----------------
    doc.add_heading("Expense Details", level=2)

    expense_data = [
        ["Sr. No.", "Expense Type", "Amount"],
        ["1", "Diesel", f"₹ {diesel:.2f}"],
        ["2", "Toll", f"₹ {toll:.2f}"],
        ["3", "Food", f"₹ {food:.2f}"],
        ["4", "Driver Charges", f"₹ {driver:.2f}"],
    ]

    sr_no = 5

    for expense in other_expenses:
        if expense["name"]:
            expense_data.append(
                [str(sr_no), expense["name"], f"₹ {expense['amount']:.2f}"]
            )
            sr_no += 1

    expense_data.append([str(sr_no), "Total Expenses", f"₹ {total_expenses:.2f}"])

    table = doc.add_table(rows=len(expense_data), cols=3)
    for i, row in enumerate(expense_data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
    add_table_borders(table)

    # ---------------- Final Summary ----------------
    doc.add_heading("Final Summary", level=2)

    summary_data = [
        ["Sr. No.", "Particular", "Amount"],
        ["1", "Total Fare", f"₹ {total_fare:.2f}"],
        ["2", "Total Expenses", f"₹ {total_expenses:.2f}"],
        ["3", "Net Profit", f"₹ {profit:.2f}"],
        ["4", "Advance Received", f"₹ {advance:.2f}"],
        ["5", "Pending Payment", f"₹ {pending_payment:.2f}"],
    ]

    table = doc.add_table(rows=len(summary_data), cols=3)
    for i, row in enumerate(summary_data):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell
    add_table_borders(table)

    doc.add_paragraph("\nSignature: _______________________")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
